"""
Document formatting setup for GeoNarrative AI MSc Thesis.
Configures styles, fonts, spacing, and helper functions.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

def create_document():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)
    
    # Configure styles
    _setup_styles(doc)
    return doc

def _setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size, bold in [
        ('Heading 1', 16, True),
        ('Heading 2', 14, True),
        ('Heading 3', 13, True),
        ('Heading 4', 12, True),
    ]:
        s = doc.styles[level]
        s.font.name = 'Times New Roman'
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_page_break(doc):
    doc.add_page_break()

def add_paragraph(doc, text, bold=False, italic=False, alignment=None, font_size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size) if font_size else Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    return p

def add_centered_text(doc, text, font_size=12, bold=False):
    return add_paragraph(doc, text, bold=bold, font_size=font_size, 
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_figure_placeholder(doc, fig_num, caption, description=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[INSERT FIGURE {fig_num} HERE]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    if description:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f"Description: {description}")
        r2.font.size = Pt(10)
        r2.italic = True
        r2.font.color.rgb = RGBColor(128, 128, 128)
    
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {fig_num}: {caption}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.italic = True
    return cap

def add_table(doc, headers, rows, caption=None, table_num=None):
    if caption and table_num:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(f"Table {table_num}: {caption}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.bold = True
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    
    doc.add_paragraph()  # spacing after table
    return table

def add_equation(doc, equation_text, eq_num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = f"  ({eq_num})" if eq_num else ""
    run = p.add_run(f"{equation_text}{label}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    return p

def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5

def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5

def add_code_block(doc, code_text, language=""):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    p._element.get_or_add_pPr().append(shading)
    return p

def add_mermaid_placeholder(doc, diagram_title, mermaid_code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[DIAGRAM: {diagram_title}]")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    add_code_block(doc, mermaid_code)
    
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("(Render the above Mermaid code at mermaid.live or paste into draw.io)")
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(128, 128, 128)

def add_screenshot_placeholder(doc, fig_num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    
    border = doc.add_paragraph()
    border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = border.add_run("─" * 50)
    r.font.color.rgb = RGBColor(180, 180, 180)
    
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = box.add_run(f"\n📷 [INSERT SCREENSHOT HERE]\n\nFigure {fig_num}\n{caption}\n")
    r1.font.size = Pt(11)
    r1.italic = True
    r1.font.color.rgb = RGBColor(100, 100, 100)
    
    border2 = doc.add_paragraph()
    border2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = border2.add_run("─" * 50)
    r2.font.color.rgb = RGBColor(180, 180, 180)
    
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(f"Figure {fig_num}: {caption}")
    rc.font.size = Pt(10)
    rc.italic = True
